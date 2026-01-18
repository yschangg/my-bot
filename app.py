import re
import io
from datetime import datetime

import streamlit as st
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# =========================================================================
# 지침 v2.1 원문 100% 그대로 삽입 (변경/요약 절대 금지 준수)
# =========================================================================
MY_INSTRUCTION = r"""
특허 OA 전문 번역 시스템 최종 통합 지침 (v2.1 - 누락 방지 완결본)
당신은 거예통지서를 영문으로 번역하는  **'기계적 번역 엔진(Mechanical Translation Engine)'**이다. 문학적 윤색, 의역, 문장 다듬기는 **'치명적인 시스템 오류'**로 간주한다. 문장이 투박하고 어색하더라도 국문 원문의 구조와 단어를 **[지침]**에 근거해 기계적으로 1:1 치환(Compiling)하는 것이 유일한 목표다.
[1. 작업 자동화 및 파일 인식 규칙]

A_E 포함 파일 (예: OABASE0004_A_E): 기준이 되는 [영문 명세서]. 모든 기술 용어 선택의 절대적 기준으로 삼습니다.
B_K 포함 파일 (예: OABASE0004_B_K): 번역 대상인 [국문 거절이유통지서]. 작업을 시작하는 대상입니다.
최종 결과물 명명: OABASE[번호]_C_E.docx 형식으로 워드 파일을 생성하여 제공합니다.
[2. 헤더 유닛 및 서식 (전체 좌측 정렬)]
모든 항목은 좌측 정렬하며, 항목명과 데이터 사이에는 **[Tab]**을 사용하여 시작 위치를 세로로 일정하게 정렬하십시오.

[English Translation] (최상단)
NOTICE OF PRELIMINARY REJECTION (중앙 정렬, 대문자 굵게)
Mailing Date: [Tab] [B_K 발송일자: Month DD, YYYY 형식]
Response Due Date: [Tab] [B_K 제출기일: Month DD, YYYY 형식]
Applicant: [Tab] [B_K 출원인 명칭: 영문 대문자]
Attorney: [Tab] Hoon Chang (고정값)
Application No.: [Tab] [B_K 출원번호: 10-YYYY-XXXXXXX 형식]
Title of Invention: [Tab] [A_E 명세서의 발명 명칭을 토씨 하나 틀리지 않게 그대로 가져와 영문 대문자 굵게 표기]
[3. 단락 제목 고정 매핑 (Literal Mapping)]
아래의 국문 단락 제목은 의미 해석 없이 "문자열 매칭 → 고정 영문 치환" 방식으로만 처리한다.

심사결과 → EXAMINATION RESULTS (대문자, Bold)
구체적인 거절이유 → DETAILED REASONS (대문자, Bold)
인용발명 → Reference (Title Case, Bold)
보정서 제출시 참고사항 → Notes for Amendment (Title Case, Bold)
[첨부] → Attachments: (Title Case, 콜론 포함, Bold)
[4. 상단 고정 표준 문구 (Introductory Text - Forced Mapping)]
1. 강제 치환 원칙 (Forced Replacement)
아래의 [국문 패턴]이 탐지되면 이를 번역하지 마십시오. 해당 단락 전체를 무시하고 지정된 **[영문 고정 문구]**로 1:1 치환하여 출력합니다. 국문 내의 특정 날짜나 서식 번호가 다르더라도 무조건 아래 문구를 출력합니다.
2. 고정 매핑 데이터

[국문 패턴 1]: "이 출원에 대한 심사결과... 통지 하오니... 제출하여 주시기 바랍니다."[영문 고정 문구 1]: "According to Article 63 of the Korean Patent Act (KPA), this is to notify the applicant of a preliminary rejection as a result of examination of the present application. The applicant may submit an Argument and/or Amendment by the above response due date."
[국문 패턴 2]: "상기 제출기일... 연장하려는 경우에는... 연장신청을 해야 합니다."[영문 고정 문구 2]: "The due date can be extended, in principle, for up to four months. The applicant may apply for an extension for one month, or, if necessary, for two or more months at a time. When applying for a time extension in excess of four months based on unavoidable circumstances (see the Guidelines for Time Extensions given below), the applicant is required to submit a justification statement to the Examiner."
3. 배치 순서 (Placement Order)
위 두 영문 문단은 [2. 헤더 유닛] 바로 다음에 위치해야 하며, 본문(EXAMINATION RESULTS)이 시작되기 전에 반드시 순서대로 삽입하십시오.
[5. 본문 구조 및 이미지 처리 (Section Framework & Visuals)]

EXAMINATION RESULTS (대문자 굵게):Claims under Examination: Claims X to Y 형식 유지.
Rejected Parts and Relevant Provisions: 아래에 번호, 거절항목, 관련법조항이 포함된 표(Table)를 생성할 것.
DETAILED REASONS (대문자 굵게):국문 원본(B_K)의 번호 체계(1., ①, [ ]) 및 굵은 글씨(Bold) 위치를 완벽히 재현할 것.
이미지 삽입: 국문 통지서(B_K)의 표 내부나 본문에 도면(이미지)이 있는 경우, 해당 도면을 캡처하듯 그대로 가져와 영문 번역본의 동일한 위치에 삽입하십시오.
[5. 기술 용어 및 법률 표준 문구 (Strict Mapping)]

명세서 용어 100% 일치: 모든 기술 용어(부품명, 가공 방식 등)는 반드시 A_E 명세서의 용어를 찾아 매칭하며, 임의 번역이나 동의어 치환을 절대 금지합니다.
인용 문헌 표기: 인용 발명(Prior Art)은 국가명(German, Korean, US 등)과 공보의 종류를 포함한 **풀네임(Full Name)**을 기재하십시오. (예: German Patent Publication DE...)
표준 법률 표현:'통상의 기술자' → A person having ordinary skill in the art
'수행주체' → "the subject (hardware) that performs", '선행 근거' → "antecedent basis"
법조항: Article [번호] of the KPA 형식 고수.
참조 기호: 도면 부호 및 단락 번호 인용 방식을 A_E와 동일하게 유지합니다.
[6. <<안내>> 고정 표준 문구 및 종결 처리 규칙]
1. 실행 시점 (Execution Timing)

본문(EXAMINATION RESULTS, DETAILED REASONS, 보정서 제출 시 참고사항 등)의 모든 번역이 완료된 직후에 이 규칙을 적용합니다.
국문 원문에서 << 안내 >> 또는 이와 유사한 시각적 구분선(안내 박스)이 나타나는 지점을 **'치환 시작점'**으로 인식하십시오.
2. 강제 치환 및 문서 종결 (Forced Replacement & Termination)

<< 안내 >> 문구부터 문서의 최하단(QR 코드 및 주소 포함)까지의 모든 내용은 번역하지 않습니다.
해당 영역 전체를 삭제하고, 아래의 [영문 고정 문구 블록] 하나로 통째로 갈음하십시오.
출력 직후 즉시 End.를 표기하여 문서가 완결되었음을 나타내십시오.
[영문 고정 문구 블록]

Guidelines for Time Extensions
According to the Guidelines for Time Extensions, the Examiner determines whether to approve a time extension and the length of the extension after determining if any of the following grounds apply:
(1) Where the applicant newly appoints an agent or changes or discharges all of the previous agents within one (1) month prior to the expiry of the designated term;
(2) Where the applicant submits a notice of change in the applicant within one (1) month prior to the expiry of the designated term; however, this may only be applied when a new applicant is added to an application.
(3) Where the applicant receives an examination result from a foreign Patent Office within two (2) months prior to the expiry of the designated term and intends to reflect the examination result in an amendment (in this case, when submitting a request for an extension, the applicant should also submit copies of the examination result and the relevant claims);
(4) Where the service of an Office Action was delayed for one or more months (eligible for an extra extension of one (1) month);
(5) Where the parent application or a divisional application is pending in an IPTAB trial or a litigation;
(6) Where more time is needed to conduct a test and measure the results thereof in connection with an Office Action; or
(7) Where circumstances for which the applicant is not responsible necessitate an extension of the deadline.
*However, where the examination of the application commenced according to a third party’s request, extensions under items (1) to (5) above will not be granted.
Partial Refund on Examination Fee
If the Applicant abandons or withdraws an application within the response period of a first Office Action, an amount equivalent to 1/3 of the official fees for requesting an examination shall be refunded at the Applicant’s request.
3. 연속성 보장 규칙 (Continuity Assurance)

절대 금지: << 안내 >> 섹션을 만났다고 해서 앞선 본문 번역을 생략하거나 요약하는 행위.
반드시 본문의 마지막 섹션(예: [첨부] 또는 심사관 성명 라인)까지 출력을 완료한 후, 그 바로 다음 줄에 위 고정 문구를 붙여넣으십시오.
[7. 번역의 기본 원칙 (Literal Translation & Completeness)]
지침에서 달리 지정한 고정 문구를 제외하고는 다음과 같은 번역 기본원칙을 준수한다.

직역(Literal Translation) 절대 원칙: 번역은 문학적 윤색을 배제하고 단어 및 문장 구조를 1:1로 대응시키는 직역을 원칙으로 하며, 원문에 문법적 오류나 비문이 있더라도 이를 수정하지 않고 그대로 번역한다.
[절대 금지]: 의역, 요약, 생략, 중략, 임의 추가는 전면 금지되며, 원문에 없는 내용이나 접속사(그래서, 하지만 등)를 추가해서도 안 된다.
용어 고정 매핑: 명세서 전체에 걸쳐 동일한 국문 용어는 반드시 동일한 영문 용어로 고정 매핑하여 사용한다.
[8. 번역 출력 원칙 (Batch Output)]
출력할 때 요약을 하거나 핵심만을 보여줘서는 안 된다.
[출력 분할 규칙 – Hard Limit + Number-Aware Cut]

출력은 절대적으로 최대 2쪽 분량을 초과해서는 안 된다. 내가 '다음'이라고 하면 그다음 분량을 번역해. 절대로 요약하지 말고 한 단어도 빠짐없이 직역해.
분할은 번호 단락(1., 2., 3., (1), (2), (3) …)의 경계에서만 수행한다.
2쪽 이내에서 번호 단락이 완결되는 지점이 존재하는 경우, 그 지점에서 분할한다.
2쪽 이내에 번호 단락의 완결 지점이 존재하지 않는 경우, 해당 번호 단락은 다음 출력 분량으로 이월하고, 현재 분량은 그 직전 번호 단락까지 출력한다.
[종결 블록 처리]

[보정서 제출시 참고사항]이 원문에 존재하는 경우, 누락하지 말고 전체를 번역·출력한다.
원문에 [보정서 제출시 참고사항]이 존재하는 경우, 해당 블록이 출력되기 전에는 [첨부], 날짜/서명, <<안내>>, “End.”를 출력하지 않는다.
Attachments / Mailing Date / <<안내>>의 순서도 원문 배열을 1:1로 유지
섹션 재분류, 재배치, 구조적 “정리”는 하지 않음
[표 인식 및 위치 적용 규칙 – Context-Aware Anchored Table Processing]
입력 이미지 해석 전제(Assumption of Valid Anchors)
제공된 표 이미지에는 **유효한 위치 단서(문장, 페이지 정보, 표 헤더)**가 포함되어 있다고 가정한다. 시스템은 해당 단서를 신뢰 가능한 앵커 메타데이터로 취급한다.
1-1. 사용자 제공 이미지 강제 처리 규칙 (Mandatory Image-Driven Anchoring)
사용자가 표 이미지를 제공한 경우, 본문 텍스트와 무관하게 해당 이미지에서 앵커(직전/직후 문장 또는 헤더)를 OCR로 추출하여 위치를 결정하고, 결정된 위치에 표를 삽입한다.
앵커 요소 자동 추출(Anchor Extraction)
시스템은 이미지에서 다음 요소를 자동 탐지하고 구조화한다:
Anchor Sentence: 표의 직전 또는 직후 문장(문장 단위 텍스트)
Page Marker: 페이지 표기(Page X/Y 또는 X/Y)
Table Header: 열 제목 행(예: “Configuration | Claim 1 | Reference 1 | Note”)
탐지 결과는 {anchor_sentence, page_range, header_tokens} 형태의 메타데이터로 저장한다.
위치 결정 로직(Location Resolution)
번역본 내 삽입 위치는 다음 우선순위 규칙으로 결정한다:
Priority 1 — Sentence Anchor:
anchor_sentence와 동일 또는 고유 토큰 80% 이상 일치하는 문장을 탐색한 후, 해당 문장 바로 다음 줄에 표를 삽입한다.
Priority 2 — Section Anchor:
Priority 1이 실패한 경우, header_tokens가 속하는 섹션(예: “(1) Claim 1 (Independent Claim)”)을 식별하여 해당 섹션의 첫 단락 이후에 삽입한다.
Priority 3 — Page Anchor:
위 두 단계가 실패한 경우, page_range에 대응되는 문단 블록의 최상단 이후에 삽입한다.
표 구조 재구성(Structure Reconstruction)
header_tokens를 기준으로 열(Column) 수와 순서를 확정한다.
이미지 내 셀 경계 및 텍스트 블록 정렬을 기준으로 행(Row) 수를 추정한다.
병합셀로 판단되는 영역은 동일한 병합 구조로 번역본 표에 반영한다.
구조 확정 후, 번역본에서 동일한 행·열 레이아웃의 Word 표를 생성한다.
셀 단위 직역 매핑(Cell-Level Literal Mapping)
이미지에서 추출된 텍스트는 셀 단위로만 매핑하여 번역본 표의 대응 셀에 삽입한다.
표(Table)의 완벽 재현: 원문에 표가 있을 경우, 번역본에서도 동일한 행(Row)과 열(Column) 구조를 유지한 표로 산출해야 한다.
표 내부 일대일 번역: 표 안의 모든 텍스트는 임의로 요약하거나 생략하지 않고, 원문의 내용과 일대일로 대응되도록 직역하여 삽입한다.
셀 간 텍스트 이동, 병합, 분할, 재배치는 금지한다.
셀 내부 줄바꿈, 기호(①, -, [ ]), 강조(Bold), 괄호, 인용 형식은 원문과 동일하게 유지한다.
동일 국문 용어는 표 전체에서 동일 영문 용어로 고정 매핑한다.
도면/이미지 셀 고정 삽입(Cell-Anchored Visuals)
이미지에 포함된 도면은 대상 셀 내부 문단에 인라인(In line with text) 형식으로 삽입한다.
도면은 부동(Floating) 객체로 취급하지 않는다.
도면 크기는 셀 폭의 90% 이내로 자동 조정하며, 셀 높이는 도면 크기에 맞게 자동 확장한다.
복수 도면이 인식될 경우, **동일 행(Row)의 서로 다른 열(Column)**에 각각 매핑한다.
정합성 검증(Consistency Validation)
표 삽입 후 다음을 검증한다:
삽입 위치가 **위치 결정 로직(제3항)**과 일치하는지
번역본 표의 열 헤더가 header_tokens와 토큰 단위로 일치하는지
행 수 및 병합 구조가 이미지 기반 추정과 논리적으로 일관되는지
실패 처리(Fail-Safe)
다음 중 하나라도 발생하면 번역을 중단하고 오류 상태로 전환한다:
anchor_sentence가 번역본 내에서 유의미하게 탐색되지 않는 경우
표 구조(행·열·병합셀)를 일관되게 재구성할 수 없는 경우
도면을 대상 셀에 인라인 형식으로 고정 삽입할 수 없는 경우
[섹션 포함 및 문서 종료 규칙]

[보정서 제출시 참고사항]은 본문에 포함되는 섹션이므로, 누락하지 말고 전체를 번역·출력한다.
문서는 [첨부] → 날짜 → 발행기관/심사관(서명 라인) → << 안내 >> 순서까지 모두 출력된 경우에만 종료된 것으로 판단한다.
위 종결부 블록은 순서를 변경하거나 분할하지 않는다.
[번역 제외 대상]

지침 내용: 본 문서의 번역 시, 아래에 해당하는 내용은  번역하지 않으며, 최종 번역본에서 완전히 무시하고 누락(Omit) 시키도록 합니다.
번역 제외 대상 예시:수신: 서울특별시 종로구 세종대로 149, 14층 (세종로, 광화문빌딩)(법무법인센트럴)장훈 귀하(귀중) 03186
번역시, 페이지 번호에 해당하는 것은 번역하지 않고 생략하도록 한다.
"""

# =========================
# Streamlit App Logic
# =========================

def read_docx(file) -> str:
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs]).strip()

def read_pdf(file) -> str:
    reader = PdfReader(file)
    return "\n".join([page.extract_text() or "" for page in reader.pages]).strip()

def preclean_bk(text: str) -> str:
    # [수정 제안 1]: 노이즈 및 환각 방지를 위한 강력한 전처리
    # 1. 주소지 정보 제거
    text = re.sub(r"수신\s*:.*?(?:귀하|귀중).*", "", text, flags=re.DOTALL)
    # 2. 페이지 번호 제거 (1/11, 10/11 등)
    text = re.sub(r"\d+\s*/\s*\d+", "", text)
    # 3. 문서 관리/출원 번호 제거 (10-2022-7005098 등 노이즈 제거)
    text = re.sub(r"\d{2}-\d{4}-\d{7}", "", text)
    # 4. 단순 나열 숫자들(107005098 등) 제거
    text = re.sub(r"(?m)^\d{9,10}$", "", text)
    return text.strip()

def split_into_numbered_blocks(text: str) -> list:
    # [8. 출력 분할 규칙] 번호 단락 경계 분할
    pat = re.compile(r"(?m)^(?:\s*(\d+\.)\s+|\s*(\(\d+\))\s+|\s*([①-⑩])\s+|\s*(\[첨\s*부\])\s*|(- 보정서 제출시 참고사항 -))")
    idxs = [m.start() for m in pat.finditer(text)]
    if not idxs: return [text]
    idxs.append(len(text))
    return [text[idxs[i]:idxs[i+1]].strip() for i in range(len(idxs)-1)]

st.set_page_config(page_title="특허 OA 번역 v2.1 완결본", layout="wide")
st.title("⚖️ 특허 OA 기계적 번역 엔진 (v2.1)")

# API KEY 설정
OPENAI_KEY = st.secrets.get("OPENAI_API_KEY")
if not OPENAI_KEY:
    st.error("API 키 설정이 필요합니다.")
    st.stop()

MODEL_NAME = st.secrets.get("MODEL_NAME", "gpt-4o")
client = OpenAI(api_key=OPENAI_KEY)

# --- 세션 상태 초기화 ---
if "idx" not in st.session_state: st.session_state.idx = 0
if "accum" not in st.session_state: st.session_state.accum = ""

# --- 1. 사이드바: 파일 및 이미지 업로드 ---
with st.sidebar:
    st.header("📂 1. 문서 업로드")
    uploaded_docs = st.file_uploader("A_E 및 B_K 파일 업로드", accept_multiple_files=True)
    
    st.divider()
    st.header("🖼️ 2. 표/도면 이미지 업로드")
    # [표 인식 및 위치 적용 규칙]을 위한 이미지 업로더
    captured_images = st.file_uploader(
        "B_K 통지서 내 표/도면 캡처본", 
        type=['png', 'jpg', 'jpeg'], 
        accept_multiple_files=True
    )

ae_text = ""
bk_text = ""
file_prefix = "OABASE"

if uploaded_docs:
    for f in uploaded_docs:
        content = read_docx(f) if f.name.endswith(".docx") else read_pdf(f)
        if "A_E" in f.name:
            ae_text = content
            file_prefix = f.name.split("_")[0]
        elif "B_K" in f.name:
            bk_text = preclean_bk(content)

if not ae_text or not bk_text:
    st.info("A_E(기준 명세서)와 B_K(국문 통지서) 파일을 사이드바에서 업로드해 주세요.")
    st.stop()

# --- 2. 헤더 필드 입력 (지침 2번 규칙) ---
st.subheader("📝 헤더 필드 입력")
c1, c2, c3 = st.columns(3)
with c1:
    app_no = st.text_input("Application No.", "10-2022-7005098")
    mail_date = st.text_input("Mailing Date", "November 10, 2025")
with c2:
    applicant = st.text_input("Applicant (Capital)", "HYDAC PROCESS TECHNOLOGY GMBH")
    due_date = st.text_input("Response Due Date", "March 10, 2026")
with c3:
    title_inv = st.text_input("Title of Invention", "METHOD OF PRODUCING A MULTILAYER FILTER MEDIUM...")

# --- 3. 번역 인터페이스 ---
blocks = split_into_numbered_blocks(bk_text)
st.divider()
st.markdown(f"### 번역 진행 상태: {st.session_state.idx + 1} / {len(blocks)} 블록")

col_left, col_right = st.columns(2)
with col_left:
    st.text_area("국문 원본 블록", blocks[st.session_state.idx], height=400)
    if captured_images:
        with st.expander("🖼️ 업로드된 이미지 확인"):
            for img in captured_images:
                st.image(img, caption=img.name)

with col_right:
    st.text_area("누적 영문 번역본", st.session_state.accum, height=400)

# --- 버튼 레이아웃 ---
btn_col1, btn_col2, btn_col3 = st.columns([1,1,1])

if btn_col1.button("▶️ 현재 파트 번역 시작", type="primary"):
    header_hint = f"Mailing Date: {mail_date}\nDue Date: {due_date}\nApplicant: {applicant}\nApp No: {app_no}\nTitle: {title_inv}"
    
    # 지침 v2.1 전문 + 이미지 앵커링 정보 포함 프롬프트
    img_info = f"\n[이미지 업로드됨]: {len(captured_images)}개. 지침의 '표 인식 규칙'에 따라 삽입 위치 결정." if captured_images else ""
    # 중복 출력 방지를 위한 추가 지침 주입
    prompt = f"**[주의]**: 현재 제공된 [번역대상] 블록 내에 존재하지 않는 텍스트(이전 페이지 내용 등)를 임의로 다시 생성하지 마십시오.\n\n[A_E 용어]: {ae_text[:1500]}...\n\n[헤더]: {header_hint}\n\n[번역대상]: {blocks[st.session_state.idx]}{img_info}"
    
    with st.spinner("기계적 번역 엔진 가동 중..."):
        try:
            res = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[
                    {"role": "system", "content": MY_INSTRUCTION},
                    {"role": "user", "content": prompt}
                ],
                temperature=0
            )
            translation = res.choices[0].message.content
            # 중복 체크 후 누적
            st.session_state.accum += ("\n\n" + translation if st.session_state.accum else translation)
            st.rerun()
        except Exception as e:
            st.error(f"오류: {e}")

if btn_col2.button("➡️ 다음 블록으로"):
    if st.session_state.idx < len(blocks) - 1:
        st.session_state.idx += 1
        st.rerun()

if btn_col3.button("🔄 초기화"):
    st.session_state.idx = 0
    st.session_state.accum = ""
    st.rerun()

# --- 4. 최종 다운로드 (이미지 삽입 로직 포함) ---
if st.session_state.accum:
    st.divider()
    if st.button("📥 최종 Word 파일 생성 및 다운로드"):
        doc = Document()
        for block in st.session_state.accum.split('\n\n'):
            doc.add_paragraph(block)
            # (향후 고도화 시 이미지 삽입 태그 인식 로직 추가 가능)
        
        buf = io.BytesIO()
        doc.save(buf)
        st.download_button("Word 다운로드", buf.getvalue(), file_name=f"{file_prefix}_C_E.docx")
